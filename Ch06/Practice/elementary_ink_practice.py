import docx
from docx.shared import RGBColor, Pt

# get text from fake message & make each line a list item
fake_text = docx.Document("fakeMessage.docx")
fake_list = [paragraph.text for paragraph in fake_text.paragraphs]  # each place 'enter' was pressed in word will be a new item in the list

# get text from real message and make each line a list item
# real_text = docx.Document('realMessage.docx')
real_text = docx.Document('realMessageChallenge.docx')
real_list = [paragraph.text for paragraph in real_text.paragraphs if len(paragraph.text) != 0]

#load template that sets style, font, margins, etc.
doc = docx.Document('template.docx')

# add letterhead
doc.add_heading('Morland Holmes', 0)  # add_heading means add a heading sytle paragraph
subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
subtitle.alignment = 1
doc.add_heading('', 1)
doc.add_paragraph('December 17, 2015')
doc.add_paragraph('')

def set_spacing(paragraph):
    """Use docx to set line spacing between paragraphs."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    

def check_blank_lines(fake_message, real_message):
    """Check to make sure there are enough blank lines in the fake message to hide the real message."""
    real_message_lines_count = len(real_message)
    fake_message_blanks = [entry for entry in fake_message if entry == ""]
    fake_message_blanks_count = len(fake_message_blanks)
    if fake_message_blanks_count < real_message_lines_count:
        return False
    else: 
        return True


if not check_blank_lines(fake_list, real_list):
    print("Not enough blank lines in the fake message.")
    quit()


length_real = len(real_list)
count_real = 0  # index of current line inr real (hidden) message

#interleave real and fake message lines
for line in fake_list:
    if count_real < length_real and line == "":
        paragraph = doc.add_paragraph(real_list[count_real])
        paragraph_index = len(doc.paragraphs) - 1
        # set real message color to white
        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255) # make it red to test
        count_real += 1
    else:
        paragraph = doc.add_paragraph(line)
    
    set_spacing(paragraph)

doc.save('ciphertext_message_letterhead.docx')

print("Done")
